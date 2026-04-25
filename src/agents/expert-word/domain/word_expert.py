from __future__ import annotations

import json
import os
import re
import sys
from typing import Any, TypedDict, List, Optional, Dict

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from bridge_protocol import BaseExpertSkill

# ── Type Definitions ─────────────────────────────────────────────────────

class WordOperation(TypedDict, total=False):
    op: str
    text: Optional[str]
    style: Optional[str]
    level: Optional[int]
    find: Optional[str]
    replace: Optional[str]
    sectionId: Optional[str]
    image_path: Optional[str]
    width_in: Optional[float]
    font_name: Optional[str]
    size_pt: Optional[int]
    bold: Optional[bool]
    target: Optional[str]
    range: Optional[Any]
    metadata: Optional[Dict[str, Any]]

class WordPayload(TypedDict):
    input: Optional[str]
    input_path: Optional[str]
    output: str
    output_path: Optional[str]
    edits: List[WordOperation]
    office_context: Optional[Dict[str, Any]]

# ── Brand defaults ────────────────────────────────────────────────────────
NEXUS_BLUE = "008CA1"
DEFAULT_FONT = "Calibri"
SUPPORTED_WORD_SUFFIXES = {".docx"}

class WordExpertSkill(BaseExpertSkill):
    """
    Industrial Word Artisan: High-fidelity document drafting and semantic editing.
    """
    def __init__(self, docx_path: Optional[str] = None, office_context: Optional[Dict[str, Any]] = None) -> None:
        self.input_path = docx_path or ""
        self.office_context = office_context or {}
        self.doc = self._load_document(docx_path) if docx_path else Document()

    def _op_insert_paragraph(self, op: WordOperation) -> None:
        self.doc.add_paragraph(str(op.get("text", "")), style=op.get("style", "Normal"))

    def _op_insert_heading(self, op: WordOperation) -> None:
        self.doc.add_heading(str(op.get("text", "")), level=op.get("level", 1))

    def _op_find_replace(self, op: WordOperation) -> None:
        find_text = str(op.get("find", ""))
        replace_text = str(op.get("replace", ""))
        for paragraph in self.doc.paragraphs:
            if find_text in paragraph.text:
                paragraph.text = paragraph.text.replace(find_text, replace_text)

    def _op_replace_section(self, op: WordOperation) -> None:
        section_id = str(op.get("sectionId", ""))
        text = str(op.get("text", ""))
        for paragraph in self.doc.paragraphs:
            if section_id in paragraph.text:
                paragraph.text = paragraph.text.replace(section_id, text)
                if op.get("style"): paragraph.style = op["style"]

    def _op_apply_named_style(self, op: WordOperation) -> None:
        style_name = op.get("style", "Normal")
        target_text = op.get("target")
        for paragraph in self.doc.paragraphs:
            if target_text is None or str(target_text) in paragraph.text:
                paragraph.style = style_name

    def _op_insert_list(self, op: WordOperation) -> None:
        items = op.get("items", [])
        style = op.get("style", "List Bullet")
        for item in items:
            self.doc.add_paragraph(str(item), style=style)

    def _op_insert_table(self, op: WordOperation) -> None:
        rows, cols = op.get("rows", 1), op.get("cols", 1)
        data = op.get("data", [])
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = op.get("style", "Table Grid")
        for r_idx, row_data in enumerate(data):
            if r_idx >= rows: break
            row_cells = table.rows[r_idx].cells
            for c_idx, val in enumerate(row_data):
                if c_idx >= cols: break
                row_cells[c_idx].text = str(val)

    def _op_add_page_break(self, _op: WordOperation) -> None:
        self.doc.add_page_break()

    def _op_set_font(self, op: WordOperation) -> None:
        target = op.get("target")
        for paragraph in self.doc.paragraphs:
            if target is None or str(target) in paragraph.text:
                for run in paragraph.runs:
                    if target is None or str(target) in run.text:
                        if op.get("font_name"): run.font.name = op["font_name"]
                        if op.get("size_pt"): run.font.size = Pt(op["size_pt"])
                        if op.get("bold") is not None: run.font.bold = op["bold"]

    def _op_add_image(self, op: WordOperation) -> None:
        self.doc.add_picture(str(op.get("image_path", "")), width=Inches(op.get("width_in", 4.0)))

    def _op_get_metadata(self, _op: WordOperation) -> Dict[str, Any]:
        return {
            "paragraph_count": len(self.doc.paragraphs),
            "styles": [s.name for s in self.doc.styles if s.type == 1]
        }

    _DISPATCH = {
        "insert_paragraph": "_op_insert_paragraph",
        "insert_heading": "_op_insert_heading",
        "find_replace": "_op_find_replace",
        "replace_section": "_op_replace_section",
        "apply_named_style": "_op_apply_named_style",
        "insert_list": "_op_insert_list",
        "insert_table": "_op_insert_table",
        "add_page_break": "_op_add_page_break",
        "set_font": "_op_set_font",
        "add_image": "_op_add_image",
        "get_metadata": "_op_get_metadata",
    }

    def save_document(self, output_path: str) -> Dict[str, Any]:
        output_extension = self._get_extension(output_path)
        if output_extension not in SUPPORTED_WORD_SUFFIXES:
            raise ValueError(f"Unsupported Word output format: {output_extension or '(missing extension)'}")
        parent_dir = os.path.dirname(output_path)
        if parent_dir and not os.path.exists(parent_dir):
            os.makedirs(parent_dir, exist_ok=True)
        self.doc.save(output_path)
        return {
            "status": "success", 
            "file": output_path, 
            "paragraph_count": len(self.doc.paragraphs),
            "output_format": output_extension.lstrip("."),
            "office_context_received": bool(self.office_context),
            "template_preserved": bool(self.input_path)
        }

    def _load_document(self, docx_path: str) -> Document:
        input_extension = self._get_extension(docx_path)
        if input_extension == ".doc":
            raise ValueError("Legacy .doc files must be converted to .docx before editing")
        if input_extension not in SUPPORTED_WORD_SUFFIXES:
            raise ValueError(f"Unsupported Word input format: {input_extension or '(missing extension)'}")
        return Document(docx_path)

    def _get_extension(self, file_path: str | None) -> str:
        return os.path.splitext(file_path or "")[1].lower()

def run(payload: WordPayload) -> Dict[str, Any]:
    input_path = payload.get("input") or payload.get("input_path")
    output_path = payload.get("output") or payload.get("output_path", "output.docx")
    expert = WordExpertSkill(input_path, payload.get("office_context"))
    applied = expert.apply_ops(payload.get("edits", []))
    for op in applied:
        if op.get("status") == "error":
            raise ValueError(op.get("message"))
    result = expert.save_document(output_path)
    result["applied_operations"] = applied
    return result

if __name__ == "__main__":
    try:
        data = json.load(sys.stdin)
        print(json.dumps(run(data)))
    except Exception as e:
        print(json.dumps({"error": str(e)}), file=sys.stderr)
        sys.exit(1)
